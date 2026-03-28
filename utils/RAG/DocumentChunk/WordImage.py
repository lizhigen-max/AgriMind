#!/usr/bin/env python
# -*- coding: UTF-8 -*-
'''
@Project ：AgriMind
@File    ：WordImage.py
@IDE     ：PyCharm
@Author  ：lizhigen
@Email   ：lizhigen1996@aliyun.com
@Date    ：2026/3/28 21:08
'''

import re
import os
from datetime import datetime
from PIL import Image
import hashlib
from .PyOCR import ocr_and_organize
import shutil
from pathlib import Path
import json
import logging
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def word_extract_and_replace_images(docx_path, output_dir="./data/extracted_content", is_ocr=False,
                                    normalize_whitespace=True, clean_formatting=True,
                                    convert_headings=True):
    """
    提取Word中的图片，并用标记文本替换图片位置，同时进行文档清理和规范化

    Args:
        docx_path: Word文档路径
        output_dir: 输出目录
        is_ocr: 是否进行OCR识别
        normalize_whitespace: 规范化全角/半角空格、制表符、连续空行
        clean_formatting: 清理多余样式，仅保留加粗、斜体
        convert_headings: 将标题样式转换为Markdown格式

    Returns:
        tuple: (processed_images列表, 新文档路径)
    """

    doc = Document(docx_path)
    os.makedirs(output_dir, exist_ok=True)

    # 创建图片保存目录
    images_dir = os.path.join(output_dir, "images")
    os.makedirs(images_dir, exist_ok=True)

    # ========== 新增：文本规范化函数 ==========
    def normalize_text(text):
        """规范化文本：统一空格、去除多余空白"""
        if not text:
            return text
        # 全角空格转半角空格
        text = text.replace('\u3000', ' ')
        # 制表符转空格
        text = text.replace('\t', ' ')
        # 多种换行符统一为\n
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        # 去除行首行尾空格
        text = '\n'.join(line.strip() for line in text.split('\n'))
        # 合并连续多个空格为单个空格
        text = re.sub(r' +', ' ', text)
        # 合并连续3个以上换行为2个换行（保留段落间隔）
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    # ========== 新增：清理格式函数 ==========
    def clean_run_formatting(run, preserve_bold=True, preserve_italic=True):
        """清理run的多余格式，仅保留指定的强调样式"""
        font = run.font

        # 清除颜色（设为自动/黑色）
        if font.color:
            font.color.rgb = None  # 设为自动颜色

        # 清除背景色/高亮
        if hasattr(font, 'highlight_color'):
            font.highlight_color = None

        # 清除字体名称（使用默认字体）
        if font.name:
            font.name = None

        # 清除字号（使用默认大小）
        if font.size:
            font.size = None

        # 清除下划线（除非需要保留）
        if font.underline:
            font.underline = False

        # 清除删除线
        if font.strike:
            font.strike = False

        # 清除阴影、轮廓等效果
        if hasattr(font, 'shadow'):
            font.shadow = False
        if hasattr(font, 'outline'):
            font.outline = False

        # 清除上标/下标
        if hasattr(font, 'superscript'):
            font.superscript = False
        if hasattr(font, 'subscript'):
            font.subscript = False

        # 保留或清除加粗
        if not preserve_bold:
            font.bold = False

        # 保留或清除斜体
        if not preserve_italic:
            font.italic = False

    # ========== 新增：标题转Markdown函数 ==========
    def convert_heading_to_markdown(paragraph):
        """将Word标题样式转换为Markdown标题格式，完全替换原有内容"""
        style_name = paragraph.style.name if paragraph.style else ""

        # 判断是否为标题样式（支持中英文命名）
        heading_level = None
        if style_name.startswith('Heading '):
            try:
                heading_level = int(style_name.split()[-1])
            except ValueError:
                pass
        elif style_name.startswith('标题'):
            try:
                heading_level = int(style_name.replace('标题', ''))
            except ValueError:
                pass

        if heading_level and 1 <= heading_level <= 9:
            # 获取纯文本内容（去除可能的列表标记）
            text_content = paragraph.text.strip()

            # 去除常见的列表前缀（如"一、", "1.", "1.1", "•", "■"等）
            import re
            # 去除编号前缀：数字、汉字数字、项目符号等
            text_content = re.sub(r'^[\\d一二三四五六七八九十]+[、.．\\s]+', '', text_content)
            text_content = re.sub(r'^[•■◆▪\\-\\*]+\\s*', '', text_content)
            text_content = text_content.strip()

            if text_content:
                # 生成Markdown标题
                markdown_heading = '#' * heading_level + ' ' + text_content

                # 完全清空段落（包括所有runs和XML子元素）
                paragraph._element.clear()

                # 重新创建段落属性（避免继承原段落的列表格式）
                from docx.oxml import OxmlElement
                pPr = OxmlElement('w:pPr')
                paragraph._element.append(pPr)

                # 添加新的run包含Markdown标题
                new_run = paragraph.add_run(markdown_heading)
                # 标题保持加粗，但清除其他格式
                clean_run_formatting(new_run, preserve_bold=True, preserve_italic=False)
                new_run.font.bold = True

                # 清除段落的对齐方式（标题通常居中，转为左对齐）
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                # 清除段落的项目符号/编号格式
                paragraph.style = doc.styles['Normal']

                return True
        return False

    # ========== 新增：处理普通段落格式 ==========
    def process_paragraph_formatting(paragraph):
        """处理普通段落的格式清理"""
        # 先清理每个run的格式
        for run in paragraph.runs:
            clean_run_formatting(run, preserve_bold=True, preserve_italic=True)

        # 规范化段落文本
        if normalize_whitespace and paragraph.text:
            full_text = paragraph.text
            normalized = normalize_text(full_text)

            if full_text != normalized:
                # 保存原有的格式标记（哪些部分是粗体/斜体）
                format_info = []
                current_pos = 0
                for run in paragraph.runs:
                    text = run.text
                    if text:
                        format_info.append({
                            'text': text,
                            'bold': run.bold,
                            'italic': run.italic,
                            'start': current_pos,
                            'end': current_pos + len(text)
                        })
                        current_pos += len(text)

                # 清空所有runs
                for run in paragraph.runs:
                    run.clear()

                # 重新添加规范化后的文本（简化处理：统一格式）
                has_bold = any(info['bold'] for info in format_info)
                has_italic = any(info['italic'] for info in format_info)

                new_run = paragraph.add_run(normalized)
                clean_run_formatting(new_run, preserve_bold=has_bold, preserve_italic=has_italic)
                if has_bold:
                    new_run.font.bold = True
                if has_italic:
                    new_run.font.italic = True

    # ========== 第一阶段：预处理所有段落（清理格式和转换标题） ==========
    if clean_formatting or convert_headings or normalize_whitespace:
        logging.info("🧹 正在清理文档格式...")
        for paragraph in doc.paragraphs:
            # 尝试转换为Markdown标题
            is_heading = False
            if convert_headings:
                is_heading = convert_heading_to_markdown(paragraph)

            # 如果不是标题，进行常规格式清理
            if not is_heading and (clean_formatting or normalize_whitespace):
                process_paragraph_formatting(paragraph)

        logging.info("✅ 格式清理完成")

    # ========== 第二阶段：提取图片（原有逻辑） ==========
    image_info_list = []
    image_count = 0

    # 先扫描并记录所有图片的位置信息
    for i, paragraph in enumerate(doc.paragraphs):
        blips = paragraph._element.xpath('.//a:blip')
        for blip in blips:
            r_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if r_id and r_id in doc.part.related_parts:
                image_count += 1
                image_info_list.append({
                    'paragraph_index': i,
                    'image_id': r_id
                })

    # 从后向前处理，避免索引变化
    image_info_list.sort(key=lambda x: x['paragraph_index'], reverse=True)
    processed_images = []

    for img_info in image_info_list:
        paragraph_index = img_info['paragraph_index']
        paragraph = doc.paragraphs[paragraph_index]

        # 提取图片
        for blip in paragraph._element.xpath('.//a:blip'):
            r_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if r_id == img_info['image_id'] and r_id in doc.part.related_parts:
                image_count = len(processed_images) + 1

                # 获取上下文文本（使用规范化后的文本）
                current_text = normalize_text(paragraph.text) if normalize_whitespace else paragraph.text
                prev_text = ""
                next_text = ""

                if paragraph_index > 0:
                    prev_text = doc.paragraphs[paragraph_index - 1].text
                    if normalize_whitespace:
                        prev_text = normalize_text(prev_text)

                if paragraph_index < len(doc.paragraphs) - 1:
                    next_text = doc.paragraphs[paragraph_index + 1].text
                    if normalize_whitespace:
                        next_text = normalize_text(next_text)

                # 提取图片并保存
                image_part = doc.part.related_parts[r_id]

                # 生成图片ID和名称
                docx_name = os.path.splitext(os.path.basename(docx_path))[0]
                docx_name_hash = hashlib.md5(docx_name.encode()).hexdigest()
                context_hash = hashlib.md5(paragraph.text.encode()).hexdigest()[:20]

                connect_text = prev_text if prev_text else next_text
                connect_text_hash = hashlib.md5(connect_text.encode()).hexdigest()
                img_id = 'img{}_{}{}_{}{}'.format(image_count, docx_name_hash[:4], docx_name_hash[-4:],
                                                      connect_text_hash[:4], connect_text_hash[-4:])
                img_name = img_id + '.png'
                img_path = os.path.join(images_dir, img_name)

                with open(img_path, 'wb') as f:
                    f.write(image_part.blob)

                # 记录图片信息
                img_data = {
                    'image_path': os.path.abspath(img_path),
                    'image_name': img_name,
                    'paragraph_index': paragraph_index + 1,
                    'original_text': current_text,
                    'context_hash': context_hash,
                    'image_id': img_id
                }
                processed_images.append(img_data)

                # OCR处理
                image_context = ''
                if is_ocr:
                    try:
                        image_context = ocr_and_organize(image_path=img_path)
                    except:
                        pass
                if len(image_context) < 100:
                    image_context = ''
                img_data.setdefault('image_context', image_context)

                # 清除段落中的所有内容
                for run in paragraph.runs:
                    run.clear()

                # 添加标记文本（简化格式，不使用醒目颜色）
                mark_text = f"【图片占位符:{img_data['image_id']}:{img_name}】"
                if image_context:
                    mark_text += f'（图片OCR内容：{image_context}）'

                run = paragraph.add_run(mark_text)

                # 应用清理后的格式：仅保留加粗用于强调
                clean_run_formatting(run, preserve_bold=True, preserve_italic=False)
                run.font.bold = True

                # 在标记文本后添加上下文信息（灰色斜体）
                if current_text.strip():
                    context_run = paragraph.add_run(f"\n[原内容]: {current_text}")
                    # 上下文信息使用更克制的格式
                    clean_run_formatting(context_run, preserve_bold=False, preserve_italic=True)
                    context_run.font.italic = True

                # 记录完整的上下文信息
                img_data['full_context'] = f"{prev_text}\n\n【图片位置】\n{current_text}\n\n{next_text}"
                img_data['previous_paragraph'] = prev_text
                img_data['next_paragraph'] = next_text

                # 保存图片元数据
                json_name = os.path.splitext(img_data['image_name'])[0] + ".json"
                json_path = os.path.join(output_dir, json_name)
                with open(json_path, 'w', encoding='utf-8') as f:
                    f.write(json.dumps(img_data, ensure_ascii=False, indent=4))

                logging.info(f"✅ 提取图片 {image_count}: {img_name}")
                if current_text:
                    logging.info(f"   上下文: {current_text[:80]}...")
                logging.info("-" * 60)

                break

    # 保存修改后的Word文档
    # src_path = Path(docx_path)
    # current_dir = Path.cwd()
    # target_dir = current_dir / "old_document"
    # target_dir.mkdir(exist_ok=True)
    #
    # dst_path = target_dir / src_path.name
    # shutil.move(str(src_path), str(dst_path))
    # doc.save(docx_path)

    new_docx_name = os.path.splitext(os.path.basename(docx_path))[0] + "_processed.docx"
    new_docx_path = os.path.join(output_dir, new_docx_name)
    doc.save(new_docx_path)

    # 保存处理摘要
    save_processing_summary(processed_images, docx_path, output_dir, new_docx_path)

    return processed_images, new_docx_path


def save_image_metadata(img_data, output_dir):
    """保存图片元数据到Markdown文件"""
    md_name = os.path.splitext(img_data['image_name'])[0] + ".md"
    md_path = os.path.join(output_dir, md_name)

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(f"# 图片元数据: {img_data['image_name']}\n\n")
        f.write(f"**图片ID**: {img_data['image_id']}\n")
        f.write(f"**所在段落**: 第{img_data['paragraph_index']}段\n")
        f.write(f"**保存路径**: {img_data['image_path']}\n")
        f.write(f"**上下文哈希**: {img_data['context_hash']}\n\n")

        f.write("## 原文内容\n")
        f.write(f"{img_data['original_text']}\n\n")

        f.write("## 图片内容\n")
        f.write(f"{img_data['image_context']}\n\n")

        f.write("## 前后文\n")
        f.write(f"**前一段**: {img_data.get('previous_paragraph', '')}\n\n")
        f.write(f"**后一段**: {img_data.get('next_paragraph', '')}\n\n")

        f.write("## 完整上下文\n")
        f.write(f"{img_data.get('full_context', '')}\n\n")

        f.write("## 图片预览\n")
        f.write(f"![{img_data['image_name']}](images/{img_data['image_name']})\n")


def save_processing_summary(processed_images, original_path, output_dir, new_docx_path):
    """保存处理摘要"""
    summary_path = os.path.join(output_dir, os.path.splitext(os.path.basename(original_path))[0] + "_处理摘要.md")

    with open(summary_path, 'w', encoding='utf-8') as f:
        f.write(f"# Word图片处理摘要\n\n")
        f.write(f"**原文件**: {os.path.basename(original_path)}\n")
        f.write(f"**处理后文件**: {os.path.basename(new_docx_path)}\n")
        f.write(f"**处理时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"**提取图片数**: {len(processed_images)} 张\n")
        f.write(f"**图片保存目录**: images/\n\n")

        f.write("## 处理流程说明\n")
        f.write("1. 从原Word中提取所有图片，保存到images目录\n")
        f.write("2. 用图片标记文本替换原图片位置\n")
        f.write("3. 保存修改后的Word文档\n")
        f.write("4. 每张图片生成对应的元数据文件(.json)\n\n")

        f.write("## 图片标记格式\n")
        f.write("`【图片占位符:img_id.png】(图片OCR内容)`\n\n")

        f.write("## 图片列表\n")
        for img in processed_images:
            f.write(f"### {img['image_id']}\n")
            f.write(f"- **文件名**: {img['image_name']}\n")
            f.write(f"- **位置**: 第{img['paragraph_index']}段\n")
            f.write(f"- **上下文哈希**: {img['context_hash']}\n")
            f.write(f"- [查看详情]({os.path.splitext(img['image_name'])[0]}.md)\n")
            f.write(f"- [图片文件](images/{img['image_name']})\n\n")


if __name__ == '__main__':
    # word_extract_and_replace_images('../Document/葡萄修枝剪型图文说明2.docx', is_ocr=True)
    word_extract_and_replace_images('../Document/技术研究中心材料汇总.docx', output_dir='./test')
